'''
Docx to text based on docx (python-docx)
'''
import os
import logging
import docx

from .logging_progress import logging_progress as progress

from nose.tools import (eq_, with_setup)

LOGGER = logging.getLogger(__name__)
LOGGER.addHandler(logging.NullHandler())


def docx_to_txt(filepath):
    '''Convert docx to text based on python-docx.'''
    if isinstance(filepath, str):
        if not os.path.exists(filepath):
            LOGGER.warning(" File %s does not exist...", filepath)
            return None

    try:
        docx_ = docx.Document(filepath)
    except Exception as exc:
        LOGGER.error(" Cannot open the file: %s", exc)
        return None

    text = ''

    totparas = len(docx_.paragraphs)
    LOGGER.info("tot: %s", totparas)
    for idx, elm in enumerate(docx_.paragraphs):
        # print(elm.text)
        progress(idx, totparas)
        text += elm.text + '\n'

    # get text frm tables if any
    for elm in docx_.tables:
        # print(elm.text)

        totrows = len(elm.rows)
        LOGGER.info("tot: %s", totrows)
        for idx, row in enumerate(elm.rows):
            progress(idx, totrows)
            rowtext = ''
            for cell in row.cells:
                rowtext += cell.text + '\t'
            text += rowtext.strip() + '\n'
        text += '\n'

    return text.strip()


def my_setup():
    fmt = '%(name)s-%(filename)s[ln:%(lineno)d]:'
    fmt += '%(levelname)s: %(message)s'
    logging.basicConfig(format=fmt, level=logging.DEBUG)


# pep8/flake8/pyling filename
# nosetests -v --nologcapture
@with_setup(my_setup)
def test_():
    ''' docx==> '''
    file = r'D:\dl\Dropbox\yeeyan\books\it-articles\2013-sep-it\2016-10-04_dual.docx'
    text = docx_to_txt(file)
    eq_(3701, len(text))


@with_setup(my_setup)
def test_doc():
    ''' doc==> '''
    file = 'D:\\dl\\Dropbox\\yeeyan\\books\\it-articles\\2013-sep-it\\2013-12-20-Oracle-dual-zh.doc'
    text = docx_to_txt(file)
    eq_(None, text)
